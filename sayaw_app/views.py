
from django.shortcuts import render,redirect,get_object_or_404,HttpResponse,HttpResponseRedirect
from django.contrib.auth import authenticate, login, logout
from .models import CustomUser,rooms,cottages,booking,paid,Guest_list,nationality,information,message_storage_1,message_storage_2,message_storage_3,gcash,paymaya
from .forms import RegistrationForm,loginForm
from django.contrib.auth.decorators import login_required, permission_required
from django.urls import reverse
from django.utils import timezone
from django.http import HttpResponse
from django.db.models import Q,Sum
import random, string
from django.http import Http404
from django.contrib import messages
from docx import Document
from django.http import HttpResponse
from io import BytesIO



def mainpage(request):
  
     return render(request,'pages/mainpage.html')


def logout_view(request):
     logout(request)
     return redirect('login') 
  
def login_process(request):
     form = loginForm()
     if request.method == "POST":
         form = loginForm(request, request.POST)
         if form.is_valid():
             username = form.cleaned_data['username']
             password = form.cleaned_data['password']
             user = authenticate(request, username=username, password=password)
             if user is not None:
                 login(request, user)
                 return redirect(reverse('reservation_App:admin_dashboard'))
             else:
                 return render(request, 'user/login.html', {"form": form,'err':'Invalid Account'})
   
     return render(request, 'pages/login.html', {"form": form})


 
     
    


                


def add_new_records(request):
     return render(request,'pages/add_new_records.html')

@login_required(login_url='reservation_App:login_process')
def admin_dashboard(request):
     user=request.user
     return render(request,'dashboard.html',{"user":user})
          


def create_account(request):
     form = RegistrationForm(request.POST)
     if request.method == "POST":
         if form.is_valid():
             username = form.cleaned_data['username']
             password = form.cleaned_data['password']
             confirm_password = form.cleaned_data['confirm_password']
             is_superadmin=form.cleaned_data['superadmin']
            
             if password == confirm_password:
                 user = CustomUser.objects.create_user(
                     username=username,
                     password=password,
                     staff=not is_superadmin, 
                     superadmin=is_superadmin
                 )
                 login(request, user)
                 return render(request,'pages/create_account.html',{'msg':'Account created successfully!'})
             else:
                 form = RegistrationForm()
                 return render(request, 'pages/create_account.html', {"form": form})

     return render(request, 'user/create_account.html', {'form': form})


def manage_establisment_view(request):
     remove=False
     add=False
     update=False
     establishment=False
     data_load=False
     phone_number=False
     if 'remove' in request.POST:
         remove=True
     if 'add' in request.POST:
         add=True
     if 'update' in request.POST:
         update=True
     if 'establishment' in request.POST:
         establishment=True
     if 'data_load' in request.POST:
         data_load=True
     if 'phone_number' in request.POST:
         phone_number=True
     context={'remove':remove,'add':add,'update':update,'establishment':establishment,'data_load':data_load,'phone_number':phone_number}
     return render(request,'pages/manage_Establishment.html',context)

def add_room(request):
     if request.method=="POST":
         is_availlable='True'
         gate_number=request.POST.get('gate_number')
         room_number=request.POST.get('room_number')
         room_price=request.POST.get('room_price')
         photo1=request.FILES.get('p1')
         photo2=request.FILES.get('p2')
         photo3=request.FILES.get('p3')
         photo4=request.FILES.get('p4')
         description=request.POST.get('description')
         room_success_message="New room added succesfully"
         add_photo=rooms(is_availlable=is_availlable,
                         gate_number=gate_number,
                         room_number=room_number,
                         room_price=room_price,
                         front_view=photo1,
                         inside_view1=photo2,
                         inside_view2=photo3,
                         inside_view3=photo4,
                         room_description=description)
         add_photo.save()
     return render(request,'pages/manage_Establishment.html',{'room_success_message':room_success_message})

def add_cottage(request):
     if request.method=="POST":
         is_availlable='True'
         gate_number=request.POST.get('gate_number')
         cottage_number=request.POST.get('cottage_number')
         cottage_price=request.POST.get('cottage_price')
         photo1=request.FILES.get('p1')
         photo2=request.FILES.get('p2')
         photo3=request.FILES.get('p3')
         photo4=request.FILES.get('p4')
         description=request.POST.get('description')
         cottage_success_message="New cottage added succesfully"
         add_photo=cottages(is_availlable=is_availlable,
                            gate_number=gate_number,
                            cottage_number=cottage_number,
                            cottage_price=cottage_price,
                            front_view=photo1,
                            inside_view1=photo2,
                            inside_view2=photo3,
                            inside_view3=photo4,
                            cottage_description=description)
         add_photo.save()
     return render(request,'pages/manage_Establishment.html',{'cottage_success_message':cottage_success_message})



def view_bookings(request):
     records=booking.objects.all()
     if request.method=="GET":
         inputValue=request.GET.get('input_value','')
         related_records=booking.objects.filter(guest_Lastname__icontains=inputValue)
         return render(request,'pages/view_bookings.html',{'query':related_records,'input':inputValue})
     return render(request,'pages/view_bookings.html',{'records':records})

def view_photos(request,id):
     room_view=get_object_or_404(rooms,id=id)

     return render(request,'pages/view_sayaw_rooms.html',{'room_view':room_view})

def view_sayaw_cottage(request,id):
     cottage=get_object_or_404(cottages,id=id)
     return render(request,'pages/view_sayaw_cottage.html',{'cottage':cottage})


    

def room_booking_view(request,id):
     cottage=cottages.objects.all()
     current_date = timezone.now()
     room=get_object_or_404(rooms,id=id)
     return render(request,'pages/reservation_form.html',{'room':room,'cottages':cottage})



    
def cottage_booking_view(request,id):
     room=rooms.objects.all()
     cottage=get_object_or_404(cottages,id=id)  
     return render(request,'pages/cottage_reservationform.html',{'cottage':cottage,'rooms':room})


def cottage_booking_process(request):
       room_price = 0
       error_msg=""
       success_msg="Thank you for visting sayaw resort"
       current_date = timezone.now()
       date_Reserved = current_date.strftime("%Y-%m-%d")
       current_time = timezone.localtime(timezone.now()).time()
       time_Reserved = current_time.strftime('%H:%M:%S')
       entrance_Fee=40
     
       if request.method=='POST':
       
         length=6
         characters = string.ascii_letters + string.digits
         transaction_code=''.join(random.choice(characters) for i in range(length))
        
         firstName=request.POST.get('firstname').capitalize()
         lastName=request.POST.get('lastname').capitalize()
         address=request.POST.get('address').capitalize()
         guest_No=request.POST.get('guest_no')
      
         cottage_price=request.POST.get('cottage_price')
         room=request.POST.get('room')
         cottage=request.POST.get('cottage_no')
         gate_No=request.POST.get('gate_no')
         date_In=request.POST.get('date_in')
         date_Out=request.POST.get('date_out')
         duration=request.POST.get('duration')
         time_In=request.POST.get('time_in')
         time_Out=request.POST.get('time_out')
         payment_method=request.POST.get('payment_method')
         cottage_price=int(cottage_price)


        # Get form data and capitalize names and address
    
        
         if cottage != '':
             cottage=int(cottage)
             reserve=get_object_or_404(cottages, cottage_number=cottage)
             reserve.is_availlable=False
             reserve.save()
         else:
             pass

         if room == '':
             room=0
         else:
             room=int(room)
         if guest_No !='':
             guest_No=int(guest_No)
         else:
             guest_No=0
      

        
         if room == '10':
             if guest_No > 10 and guest_No < 15:
                     room_price = 7000
             elif guest_No > 15 :
                  room_price = 10000
         else:       
             try:
                 room=int(room)
                 reserved=get_object_or_404(rooms,room_number=room)
                 room_price=reserved.room_price
                 reserved.is_availlable=False
                 reserved.save()
             except:
                     room_price=0
               
      
            
    
        
         if duration!=0 :
             cottage_price=int(cottage_price)*int(duration)
        
         total_bill=(entrance_Fee*int(guest_No))+(int(room_price+cottage_price))
        
            
        
         if firstName !='' and lastName !='' and address !=''  and guest_No !='' and date_In !='' and date_Out != '' and duration != '' and time_In != '' and time_Out !='':
             collecting_data=booking(payment_method=payment_method,
                                     date_Reserved=date_Reserved,
                                     time_Reserved=time_Reserved,
                                     transaction_Code=transaction_code,
                                     guest_Name=firstName,
                                     guest_Lastname=lastName,
                                     guest_Address=address,
                                     guest_Number=guest_No,
                                     cottage_number=cottage,
                                     cottage_price=cottage_price,
                                     room_number=room,
                                     gate_Number=gate_No,
                                     room_Price=room_price,
                                     date_In=date_In,
                                     date_Out=date_Out,
                                     stay_Duration=duration,
                                     time_In=time_In,
                                     time_Out=time_Out,
                                     total_bill=total_bill)
             collecting_data.save()
             return render(request,'pages/cottage_reservationform.html',{'transaction_code':transaction_code,'success_msg':success_msg})
          
            
         else:
             error_msg="There was an error while processing your booking. Please try again later"

             return render(request,'pages/reservation_form.html',{'error_msg':error_msg})
            
            
       return render(request,'pages/reservation_form.html')
   

def room_booking_process(request):
    error_msg=""
    success_msg="Thank you for visting sayaw resort"
    current_date = timezone.now()
    date_Reserved = current_date.strftime("%Y-%m-%d")
    current_time = timezone.localtime(timezone.now()).time()
    time_Reserved = current_time.strftime('%H:%M:%S')
    entrance_Fee=40
    cottage_price=0
    if request.method=='POST':
       
        length=6
        characters = string.ascii_letters + string.digits
        transaction_code=''.join(random.choice(characters) for i in range(length))
        
        firstName=request.POST.get('firstname').capitalize()
        lastName=request.POST.get('lastname').capitalize()
        address=request.POST.get('address').capitalize()
        guest_No=request.POST.get('guest_no')
        cottage=request.POST.get('cottage')
        room_No=request.POST.get('room_no')
        gate_No=request.POST.get('gate_no')
        room_price=request.POST.get('room_price')
        date_In=request.POST.get('date_in')
        date_Out=request.POST.get('date_out')
        duration=request.POST.get('duration')
        time_In=request.POST.get('time_in')
        time_Out=request.POST.get('time_out')
        payment_method=request.POST.get('payment_method')

        guest_No=int(guest_No)
        duration=int(duration)
       
        room_price=int(room_price)
        
        if guest_No > 10 and guest_No <= 15:
             room_price = 7000
         
        elif guest_No > 15 :
             room_price = 10000 
      
        
        if room_No !='' or room_No != 0:
            room=get_object_or_404(rooms,room_number=room_No)
            room.is_availlable=False
            room.save()
               
        
        if cottage != '':
             try:
                 cottage=int(cottage)
                
                 reserved=get_object_or_404(cottages,cottage_number=cottage)
                 reserved.is_availlable=False
                 cottage_price=reserved.cottage_price
                 reserved.save()
             except Http404:
                
                    cottage_price=0
            
        else:
              cottage_price=0
              cottage=0
            
     
              if duration!=0:
                    room_price=room_price*duration
        
                    total_bill=(entrance_Fee*guest_No)+room_price+cottage_price
                    # total_bill=guest_No*entrance_Fee+(room_price+table_price+cottage_price)
                    if firstName !='' and lastName !='' and address !=''  and guest_No !='' and date_In !='' and date_Out != '' and duration != '' and time_In != '' and time_Out !='':
                        collecting_data=booking(payment_method=payment_method,
                                     date_Reserved=date_Reserved,
                                     time_Reserved=time_Reserved,
                                     transaction_Code=transaction_code,
                                     guest_Name=firstName,
                                     guest_Lastname=lastName,
                                     guest_Address=address,
                                     guest_Number=guest_No,
                                     cottage_number=cottage,
                                     cottage_price=cottage_price,
                                     room_number=room_No,
                                     gate_Number=gate_No,
                                     room_Price=room_price,
                                     date_In=date_In,
                                     date_Out=date_Out,
                                     stay_Duration=duration,
                                     time_In=time_In,
                                     time_Out=time_Out,
                                     total_bill=total_bill)
                        collecting_data.save()
                        return render(request,'pages/reservation_form.html',{'transaction_code':transaction_code,'success_msg':success_msg})
          
            
                    else:
                        error_msg="There was an error while processing your booking. Please try again later"

                        return render(request,'pages/reservation_form.html',{'error_msg':error_msg})
            
            
    return render(request,'pages/reservation_form.html')
  
def transaction_code_checking(request):
     gcash_number=gcash.objects.all()
     paymaya_number=paymaya.objects.all()
     for service_num1 in gcash_number:
         num1=service_num1.number
     for service_num2 in paymaya_number:
         num2=service_num2.number
     if request.method=="POST":
         try:
             code=request.POST.get('transaction_code')
             guest=get_object_or_404(booking,transaction_Code=code)
             return render(request,'pages/guest_dashboard.html',{'guest':guest,'gcash':num1,'paymaya':num2})
         except Http404:
             error_message='Transaction Code does not exist'
             return render(request,'pages/transaction_code_checking.html',{'err':error_message})
    
     return render(request,'pages/transaction_code_checking.html',)


  
        
def view_guest_records(request,id):
     guest=get_object_or_404(booking, id=id)
     return render(request,'pages/view_guest_records.html',{'guest':guest})
        
        
def pay(request):
     current_date = timezone.now()
     date_Reserved = current_date.strftime("%Y-%m-%d")
     current_time = timezone.localtime(timezone.now()).time()
     time_Reserved = current_time.strftime('%H:%M:%S')

     error_msg=''
     success_msg=''
     if request.method=="POST":
         guest_name=request.POST.get('guestName')
         guest_lastname=request.POST.get('guest_lastname')
         imageUpload=request.FILES.get('imageUpload')
         transaction_code=request.POST.get('transaction_code')
         total_bill=request.POST.get('total_bill')
         if guest_name and transaction_code and imageUpload and total_bill and guest_lastname:
             money=paid(guest_name=guest_name,
                    payment_proof=imageUpload,
                    transaction_code=transaction_code,
                    date_paid=date_Reserved,
                    time_paid=time_Reserved,total_bill=total_bill,guest_lastname=guest_lastname)
             money.save()
             success_msg='Payment Sent!'
           
             return render(request,'pages/guest_dashboard.html',{'success_msg':success_msg,'guestname':guest_name,'guestlastname':guest_lastname})    
         else:
             error_msg='There was an issue while trying to send !!!'
             return render(request,'pages/guest_dashboard.html',{'error_msg':error_msg})
     return render(request,'pages/guest_dashboard.html')
            

def payments_viewing(request):
     guest_paids=paid.objects.all()
     return render(request,'pages/payment_viewing.html',{'guest_paids':guest_paids})


def view_proof(request,id):
     guest=get_object_or_404(paid,id=id)
     return render(request,'pages/view_picture.html',{'guest':guest})

def turn_paid(request):
   
     msg=''
     if request.method=="POST":
         insert_all_paids()
         guest_lastname=request.POST.get('lastname')
         id=request.POST.get('id')
         guest=get_object_or_404(booking,id=id)
         partial_amount=request.POST.get('partial_amount')
         partial_amount=int(partial_amount)
         if partial_amount != 0:
             total=int(guest.total_bill)
             fifty_percent=total/2
             balance=int(guest.balance_bill)
             paid=int(guest.partial_paid)
             paid=paid+partial_amount
             balance=total-paid
             guest.balance_bill=balance
             guest.partial_paid=paid
             if guest.balance_bill<=fifty_percent:
                 guest.is_paid=True
             if paid==total or balance <=0:
                 guest.balance_bill=0
                 guest.partial_paid=total
                 guest.is_fully_paid=True
                
             guest.save()
           
            
            
           
            
         msg='Mr./Mrs.'+''+''+guest_lastname+''+'is paid'
     return render(request,'pages/view_guest_records.html',{'msg':msg,'guest':guest,'messages':messages})



def delete_record(request,id):
     guest=get_object_or_404(booking, id=id)
     if request.method=="POST":
         try:
             guest.delete()
             messages.success(request, 'Object deleted successfully.')
             return redirect(reverse('reservation_App:view_bookings'))
         except Http404:
             messages.warning(request, "Record not found.")
     return render(request,'pages/delete.html',{'guest':guest})


def update_views(request,id):
     guest=get_object_or_404(booking,id=id)
     return render(request,'pages/update.html',{'guest':guest})
    
def update_record(request, id):
     guest = get_object_or_404(booking, id=id)

     if request.method == "POST":
         try:
             guest.guest_Name = request.POST.get('firstname')
             guest.guest_Lastname = request.POST.get('lastname')   
             guest.guest_Address = request.POST.get('address')
             guest.guest_Number = request.POST.get('guest_no')
             guest.child_Guest = request.POST.get('child_no')
             guest.cottage = request.POST.get('cottage')
             guest.table = request.POST.get('table')
             guest.room_Number = request.POST.get('room_no')
             guest.room_Price = request.POST.get('room_price')  
             guest.gate_Number = request.POST.get('gate_no')
             guest.stay_Duration = request.POST.get('duration')
             guest.save()
            
         except:
             return render(request, 'pages/update.html')
        
         return render(request, 'pages/update.html', {'guest': guest})

     return render(request, 'pages/update.html', {'guest': guest})


def add_guest(request):
     entrance_fee=40
     discount=20
     room=rooms.objects.filter(is_availlable=True)
     cottage=cottages.objects.filter(is_availlable=True)
     all_nation=nationality.objects.all()
     objects={'cottages':cottage,'rooms':room,'all_nation':all_nation,}
    
     if request.method=="POST":
         objects={'rooms':room,'cottages':cottage,'all_nation':all_nation}
         room=rooms.objects.filter(is_availlable=True)
         cottage=cottages.objects.filter(is_availlable=True)
         all_nation=nationality.objects.all()
         discount=20
         guest_nationality=request.POST.get('tourist_nationality')
         fname=request.POST.get('firstname').capitalize()
         lname=request.POST.get('lastname').capitalize()
         address=request.POST.get('address').capitalize()
         status=request.POST.get('status')
         pool=request.POST.get('pool')
         sel_cottage=request.POST.get('cottage')
         sel_room=request.POST.get('room')
         program=request.POST.get('program')
            
         if program == 'day':
                    
             if sel_room =='':
                sel_room=0
                room_price=0
             else:
                 selected_room=get_object_or_404(rooms, room_number=sel_room)
                 room_price=int(selected_room.room_price)
                 selected_room.is_availlable=False
                 selected_room.save()
                        
             if sel_cottage =='':
                 sel_cottage=0
                 cottage_price=0
             else:
                 selected_cottage=get_object_or_404(cottages,cottage_number=sel_cottage)
                 cottage_price=int(selected_cottage.cottage_price)
                 selected_cottage.is_availlable=False
                 selected_cottage.save()
             if pool == 'YES':
                 pool_fee = 100
             else: 
                 pool_fee = 0
                    
            
             if status == 'yes':
                 entrance_fee-=discount
             else: entrance_fee=40
                    
             total_bill=entrance_fee+room_price+cottage_price+pool_fee    
                    
             collecting_data=Guest_list(guest_name=fname,
                                         guest_lastname=lname,
                                         guest_address=address,
                                         cottage=sel_cottage,
                                         room=sel_room,
                                         nationality=guest_nationality,
                                         program=program,
                                         total_bill=total_bill
                                        
                                         )
             collecting_data.save()
             context={'total_bill':total_bill,'fname':fname,'lname':lname,'cottages':cottage,'rooms':room,'all_nation':all_nation,}
             return render(request,'pages/add_guest.html',context)
        
         else:
             if sel_room == '':
                sel_room = 0
                room_price=0
             else:
                 selected_room=get_object_or_404(rooms, room_number=sel_room)
                 room_price=int(selected_room.room_price)
                 selected_room.is_availlable=False
                 selected_room.save()
                        
             if sel_cottage =='':
                 sel_cottage=0
                 cottage_price=0
             else:
                 selected_cottage=get_object_or_404(cottages,cottage_number=sel_cottage)
                 cottage_price=int(selected_cottage.cottage_price)
                 selected_cottage.is_availlable=False
                 selected_cottage.save()
             if pool == 'YES':
                 pool_fee = 100
             else: 
                 pool_fee = 0
                    
            
             if status == 'yes':
                entrance_fee-=discount
             else:
                 entrance_fee=50
                    
             total_bill=entrance_fee+room_price+cottage_price+pool_fee    
                    
             collecting_data=Guest_list(guest_name=fname,
                                         guest_lastname=lname,
                                         guest_address=address,
                                         cottage=sel_cottage,
                                         room=sel_room,
                                         nationality=guest_nationality,
                                         program=program,
                                         total_bill=total_bill
                                        
                                         )
             collecting_data.save()
             context={'total_bill':total_bill,'fname':fname,'lname':lname,'cottages':cottage,'rooms':room,'all_nation':all_nation,}
             return render(request,'pages/add_guest.html',context)
            
        
                
     return render(request,'pages/add_guest.html',objects)
                
       
   

def track_records(request):
    
     if request.method=="POST":
         try:
             query=request.POST.get('query')
             guest=booking.objects.get(transaction_Code=query)
             return render(request,'pages/track_guest.html',{'guest':guest,'query':query})
         except booking.DoesNotExist:
             err="No records found on this code "
             query=request.POST.get('query')
             return render(request,'pages/track_guest.html',{'err':err,'code':query})
    
     return render(request,'pages/track_guest.html')


def view_rooms(request):
     all_rooms=rooms.objects.all()[:15]
     occupied_date=booking.objects.all()
     all_cottage=cottages.objects.all()
     return render(request,'pages/view_services.html',{"all_rooms":all_rooms,
                                            "all_cottages":all_cottage,
                                            "occupied_date":occupied_date})
    
def guest_list(request):
     guest_records=Guest_list.objects.all().order_by('guest_name')
    
     context={'guest':guest_records}
     return render(request,'pages/guest_list.html',context)

def confirmation(request):
     if request.method=="POST":
         return redirect(reverse('reservation_App:mainpage'))
    
    
def data_loading(request):
    
     if request.method=='POST':
         data=request.POST.get('data')
         process=information(data=data)
         process.save()
         return render(request,'pages/manage_Establishment.html',{'chatbot_success':'Data loaded successfully!'})
        
     
    
def chat(request):
     status=''
    
   
     convo=message_storage_1.objects.all().reverse()
     if request.method=='POST':
         message=request.POST.get('ask','')
         response=information.objects.filter(data__contains=message)
            
         if message != '':
             save_message=message_storage_1(message=message)
             save_message.save()
             status="Sent"
            
         return render(request,'pages/chat.html',{'convo':convo,'status':status,'response':response})
        
     return render(request, 'pages/chat.html',{'convo':convo})

def cancel_booking(request):
     guest=booking.objects.all()
    
     if request.method=='POST':
         transaction_code=request.POST.get('transaction_code')
         guest=get_object_or_404(booking,transaction_Code=transaction_code)
         guest.changeschedule_approve=True
         guest.save()  
     return render(request,'pages/guest_dashboard.html',{'guest':guest})

def change_schedule(request):

     guest=booking.objects.filter(change_schedule=True)
     return render(request,'pages/cancel_booking.html',{'cancel':guest})


def approve_cancellation(request):
     if request.method=="POST":
         transaction_code=request.POST.get('transaction_code')
         guest=get_object_or_404(booking, transaction_Code=transaction_code)
         if guest.change_schedule==False:
             guest.changeschedule_approve=True
             guest.save()
         elif guest.change_schedule==True:
             guest.changeschedule_approve=False
             guest.save()
         return redirect(reverse('reservation_App:cancel-list'))
    
def add_nationality(request):
     if request.method=="POST":
         msg=''
         nationality_name=request.POST.get('nationality').capitalize()
         process=nationality.objects.create(guest_nationality=nationality_name).save()
         if process != KeyError:
             msg='Nationality added successfully!'
            
         else: msg="Error encounter during processing"
         return render(request,'pages/manage_Establishment.html',{'msg':msg,'nationality':nationality_name})

def refresher(request):
     if request.method=='POST':
         return render(request,'pages/manage_Establishment.html')
     return render(request,'pages/manage_Establishment.html')
    
def remove_services(request):
     if request.method=="POST":
         services=request.POST.get('services')
         id=request.POST.get('number')
        
         if services == 'room':
             try:
                 msg="Room with number"
                 objects=get_object_or_404(rooms,room_number=id)
                 objects.delete()
                 return render(request,'pages/manage_Establishment.html',{'remove_msg':msg,'id':id})                
             except Http404: 
                           return HttpResponse('<br><br><br><center><h1 style="color:maroon;">The room you are trying to remove doest not exist.</h1></center>')
                      
            
         elif services == 'cottage':
             try:
                 msg="Cottage with number" 
                 objects=get_object_or_404(cottages,cottage_number=id)
                 objects.delete()
                 return render(request,'pages/manage_Establishment.html',{'remove_msg':msg,'id':id})                
             except Http404: 
                           return HttpResponse('<br><br><br><center><h1 style="color:maroon;">The cottage you are trying to remove doest not exist.</h1></center>')
            
         return render(request,'pages/manage_Establishment.html')
                    
def update_services(request):
     if request.method=='POST':
         services=request.POST.get('services')
         number=request.POST.get('number')
         gate_number=request.POST.get('gate_number')
         price=request.POST.get('price')
         image1=request.FILES.get('image1')
         image2=request.FILES.get('image2')
         image3=request.FILES.get('image3')
         image4=request.FILES.get('image4')
         description=request.POST.get('description')
        
        
         if services =='room':
             update_msg="Room with number "
             try:
                 guest=get_object_or_404(rooms, room_number=number)
                 if image1 is None :
                     image1=guest.front_view
                 else: image1=image1
                 if image2 is None:
                     image2=guest.inside_view1
                 else: image2=image2
                 if image3 is None:
                     image3=guest.inside_view2
                 else:image3=image3
                 if image4 is None:
                     image4=guest.inside_view3
                 else: image4=image4
                 guest.gate_number=gate_number
                 guest.room_price=price
                 guest.room_description=description
                 guest.front_view=image1
                 guest.inside_view1=image2
                 guest.inside_view2=image3
                 guest.inside_view3=image4
                 guest.save()
                
                 return render(request,'pages/manage_Establishment.html',{'room_update_msg':update_msg,'number':number})
             except Http404:
                             return HttpResponse('<br><br><br><center><h1 style="color:maroon;">The room you are trying to update doest not exist.</h1></center>')
                        
         elif services =='cottage':
             update_msg="Cottage with number "
             try:
                 guest=get_object_or_404(cottages, cottage_number=number)
                 if image1 is None :
                     image1=guest.front_view
                 else: image1=image1
                 if image2 is None:
                     image2=guest.inside_view1
                 else: image2=image2
                 if image3 is None:
                     image3=guest.inside_view2
                 else:image3=image3
                 if image4 is None:
                     image4=guest.inside_view3
                 else: image4=image4
                 guest.gate_number=gate_number
                 guest.cottage_price=price
                 guest.cottage_description=description
                 guest.front_view=image1
                 guest.inside_view1=image2
                 guest.inside_view2=image3
                 guest.inside_view3=image4
                 guest.save()
                
                 return render(request,'pages/manage_Establishment.html',{'cottage_update_msg':update_msg,'number':number})
             except Http404:
                             return HttpResponse('<br><br><br><center><h1 style="color:maroon;">The cottage you are trying to update doest not exist.</h1></center>')
       
def payment_effective(request,id):
     paid_guest=get_object_or_404(paid,id=id)  
     return render(request,'pages/payment_approval.html',{'guest':paid_guest})

def down_payment(request):
      if request.method=="POST":
         insert_all_paids()
         transaction_code=request.POST.get('transaction_code')
         amount=request.POST.get('amount')
         guest=get_object_or_404(booking,transaction_Code=transaction_code)
         try:
           
             guest.partial_paid+=int(amount)
             balance=int(guest.total_bill)-int(guest.partial_paid)
             guest.balance_bill=balance
             fifty_percent=int(guest.total_bill)/2
             guest.save()
             if guest.partial_paid >= fifty_percent:
                 guest.is_paid=True
                 guest.balance_bill=balance
                 guest.save()
               
             if guest.partial_paid >= guest.total_bill:
                 guest.is_paid=True
                 guest.is_fully_paid=True
                 guest.partial_paid=guest.total_bill
                 guest.balance_bill=0
                
                 guest.save()
             return render(request,'pages/payment_approval.html',{'info':guest,'amount':amount,})
            
         except Http404:
                         return render(request,'pages/payment_approval.html')
                    
      return render(request,'pages/payment_approval.html')
def goto_paidlist(request):
     if request.method=="POST":
         clear_feedback(request)
         return redirect(reverse('reservation_App:payments_viewing'))
    


def export_to_word(request):
     if request.method == "POST":
         data = Guest_list.objects.all()
         document = Document()
         document.add_heading("Generated report", level=1)
         header_row = document.add_paragraph()
         header_row.add_run("Firstname").bold = True
         header_row.add_run("\t\t")
         header_row.add_run("Lastname").bold = True
         header_row.add_run("\t\t")
         header_row.add_run("Address").bold = True
         header_row.add_run("\t\t")
         header_row.add_run("Nationality").bold = True
         header_row.add_run("\t\t")
       
         for guest in data:
            document.add_paragraph(f" {guest.guest_name}\t\t {guest.guest_lastname} \t\t{guest.guest_address}\t\t\t {guest.nationality}") 
     output = BytesIO()
     document.save(output)
     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
     response['Content-Disposition'] = 'attachment; filename=generated_reports.docx'
     response.write(output.getvalue())
     return response
  
    

def feedback(request,id):
     guest=get_object_or_404(paid,id=id)
        
     return render(request,'pages/feedback.html',{'guest':guest})

def feedback_process(request):
   
     if request.method=="POST":
        
         code=request.POST.get("code")
         message=request.POST.get('message')
         subject=booking.objects.all().get(transaction_Code=code)
         subject.feedback=message
         subject.save()
        
         return render(request,'pages/feedback.html',{'subject':subject,'message':'Feedback Sent'})

  
     return render(request,'pages/feedback.html')



def insert_number(request):
     if request.method== 'POST':
         gcash_number=request.POST.get("gcash")
         paymaya_number=request.POST.get("paymaya")
         paymaya.objects.update(number=paymaya_number)
         gcash.objects.update(number=gcash_number)
     return render(request,'pages/manage_Establishment.html',{"suc":"Phone number save successfully"})


def clear_feedback(request):
     if request.method=="POST":
        booking.objects.filter(is_paid = True).update(feedback="")
        return redirect(reverse('reservation_App:goto_paidlist'))
     return render(request,'pages/manage_Establishment.html')
    
        
        
def manage_account(request):
     form = RegistrationForm(request.POST)
     create=False
     if 'create' in request.POST:
         return redirect(reverse('reservation_App:create_account'))

     return render(request,'pages/manage_account.html',{'create':create,'form':form})

def insert_all_paids():
     all_paids = booking.objects.filter(is_fully_paid=True)

     for booking_obj in all_paids:
         if not Guest_list.objects.filter(guest_name=booking_obj.guest_Name, guest_lastname=booking_obj.guest_Lastname,guest_address=booking_obj.guest_Address).exists():
             collect = Guest_list(
                 guest_name=booking_obj.guest_Name,
                 guest_lastname=booking_obj.guest_Lastname,
                 guest_address=booking_obj.guest_Address,
                 nationality='Filipino',
                 total_bill=booking_obj.total_bill
             )
             collect.save()

            
def reset(request):
     form=False
     if request.method=="POST":
         form=True
     return render(request,'pages/add_guest.html',{'form':form})

def updating(request):
     if 'one' in request.POST:
         option=request.POST.get('option')
         number=request.POST.get('number')
         if option !="":
             if option == "cottage":
                 msg="Cottage with number: "
                 cottage=get_object_or_404(cottages, cottage_number=number)
                 cottage.is_availlable=True
                 cottage.save()
                 return render(request,'user/add_guest.html',{"msg":msg,'number':number})
             elif option == "room":
                 msg="Room with number: "
                 room=get_object_or_404(rooms, room_number=number)
                 room.is_availlable=True
                 room.save()
                 return render(request,'user/add_guest.html',{"msg":msg,'number':number})
             else:
                 return render(request,'user/add_guest.html')
         return render(request,'user/add_guest.html')
     elif 'all' in request.POST:
             rooms.objects.all().update(is_availlable=True)
             cottages.objects.all().update(is_availlable=True)
             return HttpResponse("<p> All services such as cottage and rooms are set to availlable</p>")
        
        
def forgot_code(request):
     code=''
     if request.method=="POST":
         firstname=request.POST.get('fname').capitalize()
         lastname=request.POST.get('lname').capitalize()
       
         if firstname != '' and lastname != '':
             try:
                 code=get_object_or_404(booking,guest_Name=firstname,guest_Lastname=lastname)
                 msg="Hey! I found your transaction code!"
                 return render(request,'pages/forgot_code.html',{'msg':msg,'code':code})
             except Http404:
                     return render(request,'pages/forgot_code.html',{'err':'Cant find you code  :(   make sure to input your information cerrectly.'})
         else:
              return render(request,'pages/forgot_code.html')
            
     return render(request,'pages/forgot_code.html')


def information_dashboard(request):
     all_guest=Guest_list.objects.all().count()
     all_bookings=booking.objects.all().count()
     total_amount = Guest_list.objects.aggregate(Sum('total_bill'))['total_bill__sum']
    
     context={"all_guest":all_guest,'total_amount':total_amount,"all_booking":all_bookings}
     return render(request,'pages/information_dashboard.html',context)


def guest_dashboard(request):
     return render(request,'user/guest_dashboard.html')


def changeSchedule(request,id):
     details=get_object_or_404(booking,id=id)
     if request.method=="POST":
       
        id=request.POST.get("id")
        reason=request.POST.get("reason")
        date=request.POST.get("date")
        time=request.POST.get("time")
        details.date_In=date
        details.time_In=time
        details.changeschedule_reason=reason
        details.change_schedule=True
     details.save()
     return render(request,'pages/changeschedule_form.html',{'info':details})

# def changing_scedule_process(request,id):
#      details=get_object_or_404(booking,id=id)
#      return render(request,'pages/changeschedule_form.html',{"info":details})


